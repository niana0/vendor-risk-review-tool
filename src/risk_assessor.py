"""
Risk Assessor - Analyzes completed questionnaire and generates risk assessment
"""
from typing import Dict, List, Any
import json


class RiskAssessor:
    """Assess risks based on questionnaire responses"""

    RISK_CATEGORIES = {
        'data_protection': ['encryption', 'data protection', 'privacy', 'gdpr', 'confidentiality'],
        'access_control': ['authentication', 'authorization', 'access control', 'mfa', 'password'],
        'monitoring': ['logging', 'monitoring', 'audit', 'siem', 'detection'],
        'incident_response': ['incident', 'response', 'breach', 'disaster recovery', 'backup'],
        'compliance': ['compliance', 'certification', 'soc 2', 'iso 27001', 'audit'],
        'vulnerability_management': ['vulnerability', 'patch', 'scanning', 'penetration test'],
        'vendor_management': ['vendor', 'third party', 'supplier', 'subprocessor']
    }

    def __init__(self):
        self.risk_assessment = {
            "risks": [],
            "recommendations": [],
            "summary": {}
        }

    def assess(self, question_mappings: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Perform risk assessment"""

        # Analyze by confidence level
        confidence_stats = self._analyze_confidence(question_mappings)

        # Identify high-risk areas
        risks = self._identify_risks(question_mappings)

        # Generate recommendations
        recommendations = self._generate_recommendations(risks, question_mappings)

        # Calculate overall risk score
        risk_score = self._calculate_risk_score(question_mappings)

        self.risk_assessment = {
            "overall_risk": risk_score['level'],
            "risk_score": risk_score['score'],
            "confidence_distribution": confidence_stats,
            "risks": risks,
            "recommendations": recommendations,
            "summary": {
                "total_questions": len(question_mappings),
                "answered_high_confidence": confidence_stats.get('HIGH', 0),
                "answered_medium_confidence": confidence_stats.get('MEDIUM', 0),
                "answered_low_confidence": confidence_stats.get('LOW', 0),
                "insufficient_evidence": confidence_stats.get('NOT_FOUND', 0),
                "critical_risks": len([r for r in risks if r['severity'] == 'HIGH']),
                "medium_risks": len([r for r in risks if r['severity'] == 'MEDIUM']),
                "low_risks": len([r for r in risks if r['severity'] == 'LOW'])
            }
        }

        return self.risk_assessment

    def _analyze_confidence(self, mappings: List[Dict[str, Any]]) -> Dict[str, int]:
        """Analyze confidence distribution"""
        stats = {}
        for mapping in mappings:
            conf = mapping.get('confidence', 'NOT_FOUND')
            stats[conf] = stats.get(conf, 0) + 1
        return stats

    def _identify_risks(self, mappings: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Identify specific risks"""
        risks = []

        # Group questions by risk category
        for category, keywords in self.RISK_CATEGORIES.items():
            category_questions = []

            for mapping in mappings:
                question_lower = mapping['question'].lower()
                if any(kw in question_lower for kw in keywords):
                    category_questions.append(mapping)

            if not category_questions:
                continue

            # Assess category risk
            low_conf_count = sum(1 for m in category_questions
                                if m['confidence'] in ['LOW', 'NOT_FOUND'])

            if low_conf_count > len(category_questions) * 0.5:
                severity = 'HIGH' if low_conf_count > len(category_questions) * 0.7 else 'MEDIUM'

                risks.append({
                    "category": category.replace('_', ' ').title(),
                    "severity": severity,
                    "description": f"Insufficient evidence for {low_conf_count}/{len(category_questions)} {category.replace('_', ' ')} controls",
                    "affected_questions": [m['question_id'] for m in category_questions if m['confidence'] in ['LOW', 'NOT_FOUND']],
                    "gaps": list(set([gap for m in category_questions for gap in m.get('gaps', [])]))
                })

        # Sort by severity
        severity_order = {'HIGH': 0, 'MEDIUM': 1, 'LOW': 2}
        risks.sort(key=lambda x: severity_order.get(x['severity'], 3))

        return risks

    def _generate_recommendations(
        self,
        risks: List[Dict[str, Any]],
        mappings: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Generate prioritized recommendations"""
        recommendations = []

        # Recommendations for high-severity risks
        for risk in risks:
            if risk['severity'] == 'HIGH':
                recommendations.append({
                    "priority": "Critical",
                    "category": risk['category'],
                    "action": f"Request additional documentation for {risk['category']} controls",
                    "rationale": risk['description'],
                    "questions_to_followup": risk['affected_questions'][:5]
                })

        # General recommendations based on confidence levels
        not_found_questions = [m for m in mappings if m['confidence'] == 'NOT_FOUND']
        if len(not_found_questions) > 5:
            recommendations.append({
                "priority": "High",
                "category": "Documentation",
                "action": "Request comprehensive security documentation package",
                "rationale": f"{len(not_found_questions)} questions have no supporting evidence",
                "questions_to_followup": [m['question_id'] for m in not_found_questions[:10]]
            })

        # Recommendations for contradictions
        # (This would require more sophisticated analysis)

        return recommendations

    def _calculate_risk_score(self, mappings: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Calculate overall risk score"""
        total = len(mappings)
        if total == 0:
            return {"score": 0, "level": "UNKNOWN"}

        # Scoring: HIGH=3, MEDIUM=2, LOW=1, NOT_FOUND=0
        score_map = {'HIGH': 3, 'MEDIUM': 2, 'LOW': 1, 'NOT_FOUND': 0}
        total_score = sum(score_map.get(m.get('confidence', 'NOT_FOUND'), 0) for m in mappings)
        normalized_score = (total_score / (total * 3)) * 100  # 0-100 scale

        # Determine risk level
        if normalized_score >= 70:
            level = "LOW RISK"
        elif normalized_score >= 50:
            level = "MEDIUM RISK"
        elif normalized_score >= 30:
            level = "HIGH RISK"
        else:
            level = "CRITICAL RISK"

        return {
            "score": round(normalized_score, 1),
            "level": level
        }

    def to_docx(self, output_path: str = "output/risk_assessment_report.docx") -> str:
        """Generate a Word document report."""
        from docx import Document
        from datetime import datetime

        doc = Document()
        doc.add_heading("Vendor Risk Assessment Report", 0)

        # Executive summary
        doc.add_heading("Executive Summary", 1)
        for label, value in [
            ("Overall Risk Level", self.risk_assessment["overall_risk"]),
            ("Risk Score", f"{self.risk_assessment['risk_score']}/100"),
            ("Assessment Date", datetime.now().strftime("%Y-%m-%d")),
        ]:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        doc.add_heading("Quick Stats", 2)
        s = self.risk_assessment["summary"]
        for label, value in [
            ("Total Questions Assessed", str(s["total_questions"])),
            ("High Confidence Answers", str(s["answered_high_confidence"])),
            ("Medium Confidence Answers", str(s["answered_medium_confidence"])),
            ("Low / No Evidence", str(s["answered_low_confidence"] + s["insufficient_evidence"])),
            ("Critical Risks Identified", str(s["critical_risks"])),
        ]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(value)

        # Key risks
        doc.add_heading("Key Risks", 1)
        for risk in self.risk_assessment["risks"]:
            doc.add_heading(f"{risk['severity']} — {risk['category']}", 2)
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(risk["description"])
            for gap in risk.get("gaps", [])[:3]:
                doc.add_paragraph(gap, style="List Bullet")

        if not self.risk_assessment["risks"]:
            doc.add_paragraph("No significant risks identified.")

        # Recommendations
        doc.add_heading("Recommendations", 1)
        for idx, rec in enumerate(self.risk_assessment["recommendations"], 1):
            doc.add_heading(f"{idx}. [{rec['priority']}] {rec['action']}", 2)
            for label, value in [
                ("Category", rec["category"]),
                ("Rationale", rec["rationale"]),
            ]:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(value)
            if rec.get("questions_to_followup"):
                p = doc.add_paragraph()
                p.add_run("Questions to Follow-up: ").bold = True
                p.add_run(", ".join(rec["questions_to_followup"][:5]))

        # Confidence distribution table
        doc.add_heading("Confidence Distribution", 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].paragraphs[0].add_run("Confidence Level").bold = True
        hdr[1].paragraphs[0].add_run("Count").bold = True
        for level, count in self.risk_assessment["confidence_distribution"].items():
            row = table.add_row().cells
            row[0].text = level
            row[1].text = str(count)

        doc.add_paragraph("")
        p = doc.add_paragraph("This report was automatically generated by the Vendor Security Assessment Tool.")
        p.runs[0].italic = True

        doc.save(output_path)
        return output_path

    def to_json(self, output_path: str = "risk_assessment.json"):
        """Save risk assessment to JSON"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.risk_assessment, f, indent=2, ensure_ascii=False)
        return output_path

    def to_markdown(self, output_path: str = "risk_assessment_report.md") -> str:
        """Generate markdown report"""
        md = f"""# Vendor Risk Assessment Report

## Executive Summary

**Overall Risk Level:** {self.risk_assessment['overall_risk']}
**Risk Score:** {self.risk_assessment['risk_score']}/100
**Assessment Date:** {self._get_date()}

### Quick Stats
- **Total Questions Assessed:** {self.risk_assessment['summary']['total_questions']}
- **High Confidence Answers:** {self.risk_assessment['summary']['answered_high_confidence']}
- **Medium Confidence Answers:** {self.risk_assessment['summary']['answered_medium_confidence']}
- **Low/No Evidence:** {self.risk_assessment['summary']['answered_low_confidence'] + self.risk_assessment['summary']['insufficient_evidence']}

---

## Key Risks

"""
        # Add risks
        for risk in self.risk_assessment['risks']:
            md += f"### {risk['severity']} - {risk['category']}\n\n"
            md += f"**Description:** {risk['description']}\n\n"
            if risk.get('gaps'):
                md += "**Evidence Gaps:**\n"
                for gap in risk['gaps'][:3]:
                    md += f"- {gap}\n"
            md += "\n"

        md += "---\n\n## Recommendations\n\n"

        for idx, rec in enumerate(self.risk_assessment['recommendations'], 1):
            md += f"### {idx}. [{rec['priority']}] {rec['action']}\n\n"
            md += f"**Category:** {rec['category']}  \n"
            md += f"**Rationale:** {rec['rationale']}\n\n"
            if rec.get('questions_to_followup'):
                md += f"**Questions to Follow-up:** {', '.join(rec['questions_to_followup'][:5])}\n\n"

        md += "---\n\n## Confidence Distribution\n\n"
        md += "| Confidence Level | Count |\n"
        md += "|-----------------|-------|\n"
        for level, count in self.risk_assessment['confidence_distribution'].items():
            md += f"| {level} | {count} |\n"

        md += "\n---\n\n"
        md += "*This report was automatically generated by the Vendor Security Assessment Tool*\n"

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)

        return output_path

    def _get_date(self):
        """Get current date"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d")
