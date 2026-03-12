"""
Shared utility: convert basic markdown text to a python-docx Document.
Handles headers (# ## ###), bold (**text**), bullets (- ), tables (| |), and paragraphs.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, doc: Document = None) -> Document:
    """Parse markdown and write into a python-docx Document. Returns the document."""
    if doc is None:
        doc = Document()
        _set_default_styles(doc)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(_strip_md(line[2:]), level=1)

        # Heading 2
        elif line.startswith("## ") and not line.startswith("### "):
            p = doc.add_heading(_strip_md(line[3:]), level=2)

        # Heading 3
        elif line.startswith("### "):
            p = doc.add_heading(_strip_md(line[4:]), level=3)

        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Table row — collect all consecutive table lines
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue  # i already advanced

        # Bullet
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", line))

        # Empty line — small spacer
        elif line.strip() == "":
            doc.add_paragraph("")

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1

    return doc


def _add_table(doc: Document, table_lines: list):
    """Render a markdown table into a docx table."""
    # Parse rows, skip separator lines (|---|---|)
    rows = []
    for line in table_lines:
        if re.match(r"^\s*\|[\s\-:]+\|", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= col_count:
                break
            cell = table.rows[r_idx].cells[c_idx]
            p = cell.paragraphs[0]
            p.clear()
            _add_inline(p, cell_text)
            # Bold the header row
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph("")  # spacing after table


def _add_inline(paragraph, text: str):
    """Add text to a paragraph, handling **bold** and *italic* inline markdown."""
    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Strip any remaining backticks
            paragraph.add_run(part.replace("`", ""))


def _strip_md(text: str) -> str:
    """Remove inline markdown from a heading string."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.replace("`", "").strip()


def _set_default_styles(doc: Document):
    """Apply sensible default font settings to a new document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
