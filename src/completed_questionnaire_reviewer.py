"""
Completed Questionnaire Reviewer - Reads vendor-answered questionnaires and flags security findings.
"""
import os
import json
from typing import Dict, List, Any
import openpyxl
from openai import OpenAI


class CompletedQuestionnaireReviewer:

    COMPLETION_MODEL = "gpt-4o"
    BATCH_SIZE = 10

    # Keywords for auto-detecting header/question/answer columns
    QUESTION_KEYWORDS = {"question", "control", "requirement", "description", "item"}
    ANSWER_KEYWORDS = {"answer", "response", "vendor response", "vendor answer", "comments"}

    def __init__(self):
        api_key = os.getenv("OPENAI_API_KEY")
        self.client = OpenAI(api_key=api_key) if api_key else None
        self.findings: List[Dict[str, Any]] = []

    # ------------------------------------------------------------------
    # Excel loading
    # ------------------------------------------------------------------

    def load_questionnaire(self, excel_path: str) -> List[Dict[str, Any]]:
        """Load Q&A pairs from a vendor-completed Excel questionnaire."""
        self._excel_path = excel_path  # stored for to_excel()
        self._header_row_idx = None
        self._question_col = None
        qa_pairs = []
        try:
            workbook = openpyxl.load_workbook(excel_path, data_only=True)
            sheet = workbook[workbook.sheetnames[0]]

            # Auto-detect header row and column indices
            header_row_idx = None
            question_col = None
            answer_col = None

            for row_idx, row in enumerate(sheet.iter_rows(min_row=1, max_row=20, values_only=True), start=1):
                if row_idx > 20:
                    break
                for col_idx, cell in enumerate(row):
                    if cell is None:
                        continue
                    cell_lower = str(cell).strip().lower()
                    if any(kw in cell_lower for kw in self.QUESTION_KEYWORDS) and question_col is None:
                        header_row_idx = row_idx
                        question_col = col_idx
                    if any(kw in cell_lower for kw in self.ANSWER_KEYWORDS) and answer_col is None:
                        answer_col = col_idx
                if header_row_idx is not None and question_col is not None and answer_col is not None:
                    break

            # Fall back to first two columns if detection fails
            if header_row_idx is None:
                header_row_idx = 1
            if question_col is None:
                question_col = 0
            if answer_col is None:
                answer_col = 1
            self._header_row_idx = header_row_idx
            self._question_col = question_col
            self._row_to_qid: Dict[int, str] = {}  # excel row number → question id

            # Read data rows — record exact row index for each question
            q_counter = 1
            for excel_row_idx, row in enumerate(
                sheet.iter_rows(min_row=header_row_idx + 1, values_only=True),
                start=header_row_idx + 1
            ):
                question = str(row[question_col]).strip() if row[question_col] is not None else ""
                answer = str(row[answer_col]).strip() if row[answer_col] is not None else ""

                if not question and not answer:
                    continue
                if not question:
                    continue

                qid = f"Q{q_counter}"
                self._row_to_qid[excel_row_idx] = qid
                qa_pairs.append({
                    "id": qid,
                    "question": question,
                    "answer": answer if answer else "No response provided",
                })
                q_counter += 1

        except Exception as e:
            print(f"Error loading questionnaire: {e}")

        return qa_pairs

    # ------------------------------------------------------------------
    # LLM review
    # ------------------------------------------------------------------

    def review(self, qa_pairs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Review Q&A pairs and flag security findings. Stores results in self.findings."""
        self.findings = []

        if not self.client or not qa_pairs:
            return self.findings

        # Process in batches
        for i in range(0, len(qa_pairs), self.BATCH_SIZE):
            batch = qa_pairs[i:i + self.BATCH_SIZE]
            batch_results = self._review_batch(batch)
            self.findings.extend(batch_results)

        return self.findings

    def _review_batch(self, batch: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Send one batch of Q&A pairs to GPT-4o and return structured results."""
        qa_text = "\n".join(
            f"[{item['id']}] Q: {item['question']}\n       A: {item['answer']}"
            for item in batch
        )
        id_to_item = {item["id"]: item for item in batch}

        prompt = (
            "You are a GRC analyst reviewing a vendor-completed security questionnaire. "
            "For each Q&A pair below, classify the vendor's answer as one of three outcomes.\n\n"
            "**Outcome 1 — Finding (is_finding: true)**\n"
            "The vendor's answer reveals a real security or compliance gap:\n"
            "- Answers No/None/N/A to having required security controls\n"
            "- Lacks a formal policy (security, privacy, acceptable use, incident response, etc.)\n"
            "- Admits to no encryption, no MFA, no penetration testing, no backups, no logging\n"
            "- Has had security breaches or incidents without adequate response\n"
            "- Uses outdated or unsupported software\n"
            "- Has no RBAC or access management\n"
            "- Has no security awareness training program\n"
            "- Has incomplete or missing incident response capability\n\n"
            "Severity guidelines:\n"
            "- High: No encryption at rest/in transit, no DR/BCP, active breach, no formal security policy, no MFA, no penetration testing, no access controls\n"
            "- Medium: Partial or informal controls, manual processes without documentation\n"
            "- Low: Minor gaps, best-practice improvements\n\n"
            "'Not applicable' may or may not be a finding — use judgment based on the question context.\n\n"
            "**Outcome 2 — Informational (is_informational: true)**\n"
            "The question asks the vendor to provide a document (e.g., SOC 2 report, penetration test report, "
            "security policy, privacy policy, certificate, audit report, or any formal document), AND the vendor's "
            "answer indicates one has been provided — e.g., 'see attached', 'see document', 'see comment', "
            "'provided separately', 'attached', 'please see', 'refer to', or similar deferral language. "
            "These are NOT findings. Mark them informational and recommend that the reviewer verify the document.\n\n"
            "**Outcome 3 — No issue (is_finding: false, is_informational: false)**\n"
            "The answer is adequate and raises no concern.\n\n"
            "Q&A pairs to review:\n"
            f"{qa_text}\n\n"
            "Respond with a JSON object with a 'results' array. Each element must have:\n"
            "- question_id (string, e.g. 'Q1')\n"
            "- is_finding (boolean)\n"
            "- is_informational (boolean) — true only when the question requests a document and the answer defers to one\n"
            "If is_finding is true, also include:\n"
            "- severity ('High', 'Medium', or 'Low')\n"
            "- finding_title (short title, e.g. 'No Information Security Policy')\n"
            "- description (2-3 sentences explaining why this is a risk)\n"
            "- recommendation (concrete action for the reviewer to take)\n"
            "If is_informational is true, also include:\n"
            "- info_title (short title, e.g. 'SOC 2 Report Referenced')\n"
            "- recommendation (e.g. 'Review the provided SOC 2 report to verify scope and audit opinion.')\n\n"
            "is_finding and is_informational are mutually exclusive — only one can be true.\n\n"
            "Example: {\"results\": [{\"question_id\": \"Q1\", \"is_finding\": false, \"is_informational\": false}, "
            "{\"question_id\": \"Q2\", \"is_finding\": true, \"is_informational\": false, \"severity\": \"High\", "
            "\"finding_title\": \"No MFA\", \"description\": \"...\", \"recommendation\": \"...\"}, "
            "{\"question_id\": \"Q3\", \"is_finding\": false, \"is_informational\": true, "
            "\"info_title\": \"SOC 2 Report Referenced\", \"recommendation\": \"Review the attached SOC 2 report.\"}]}"
        )

        try:
            response = self.client.chat.completions.create(
                model=self.COMPLETION_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
                response_format={"type": "json_object"},
            )
            data = json.loads(response.choices[0].message.content)
            raw_results = data.get("results", [])
        except Exception as e:
            print(f"Batch review failed: {e}")
            return []

        # Merge LLM results with original Q&A data
        enriched = []
        for r in raw_results:
            qid = r.get("question_id")
            original = id_to_item.get(qid, {})
            entry = {
                "question_id": qid,
                "question": original.get("question", ""),
                "vendor_answer": original.get("answer", ""),
                "is_finding": r.get("is_finding", False),
                "is_informational": r.get("is_informational", False),
            }
            if r.get("is_finding"):
                entry["severity"] = r.get("severity", "Medium")
                entry["finding_title"] = r.get("finding_title", "")
                entry["description"] = r.get("description", "")
                entry["recommendation"] = r.get("recommendation", "")
            elif r.get("is_informational"):
                entry["info_title"] = r.get("info_title", "Document Referenced")
                entry["recommendation"] = r.get("recommendation", "Review the referenced document.")
            enriched.append(entry)

        return enriched

    # ------------------------------------------------------------------
    # Output: Excel (original + findings columns)
    # ------------------------------------------------------------------

    def to_excel(self, output_path: str) -> str:
        """Write questionnaire + Finding column to a new workbook from self.findings data."""
        from openpyxl import Workbook
        from openpyxl.styles import PatternFill, Font, Alignment
        from openpyxl.utils import get_column_letter

        SEV_FILLS = {
            "High":          PatternFill("solid", fgColor="C0392B"),
            "Medium":        PatternFill("solid", fgColor="E67E22"),
            "Low":           PatternFill("solid", fgColor="27AE60"),
            "Informational": PatternFill("solid", fgColor="2980B9"),
        }
        WHITE_FONT = Font(color="FFFFFF", bold=True)
        BOLD = Font(bold=True)
        WRAP = Alignment(wrap_text=True, vertical="top")
        CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)

        wb = Workbook()
        ws = wb.active
        ws.title = "Questionnaire Review"

        # Headers
        headers = ["#", "Question", "Vendor Answer", "Finding"]
        for col, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = BOLD
            cell.alignment = CENTER

        # Data rows
        for i, entry in enumerate(self.findings, start=1):
            is_finding = entry.get("is_finding", False)
            is_info = entry.get("is_informational", False)
            sev = entry.get("severity", "") if is_finding else ""
            title = entry.get("finding_title", "") if is_finding else entry.get("info_title", "") if is_info else ""

            ws.cell(row=i + 1, column=1, value=entry.get("question_id", f"Q{i}")).alignment = CENTER
            ws.cell(row=i + 1, column=2, value=entry.get("question", "")).alignment = WRAP
            ws.cell(row=i + 1, column=3, value=entry.get("vendor_answer", "")).alignment = WRAP

            finding_cell = ws.cell(row=i + 1, column=4)
            if is_finding:
                finding_cell.value = f"{sev} — {title}" if title else sev
                finding_cell.fill = SEV_FILLS.get(sev, PatternFill())
                finding_cell.font = WHITE_FONT
                finding_cell.alignment = CENTER
            elif is_info:
                finding_cell.value = f"Informational — {title}" if title else "Informational"
                finding_cell.fill = SEV_FILLS["Informational"]
                finding_cell.font = WHITE_FONT
                finding_cell.alignment = CENTER

        # Column widths
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 50
        ws.column_dimensions["C"].width = 50
        ws.column_dimensions["D"].width = 30

        wb.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Output: docx
    # ------------------------------------------------------------------

    def to_docx(self, output_path: str) -> str:
        """Save findings to a Word document grouped by severity."""
        from docx import Document

        doc = Document()
        doc.add_heading("Completed Questionnaire — Security Findings", 0)

        findings = [f for f in self.findings if f.get("is_finding")]
        info_items = [f for f in self.findings if f.get("is_informational")]
        total = len(self.findings)

        # Summary
        summary_para = doc.add_paragraph()
        counts = {s: sum(1 for f in findings if f.get("severity") == s) for s in ["High", "Medium", "Low"]}
        summary_para.add_run(
            f"Total questions reviewed: {total}   |   Findings: {len(findings)}   |   "
            f"High: {counts['High']}  Medium: {counts['Medium']}  Low: {counts['Low']}   |   "
            f"Informational: {len(info_items)}"
        )

        doc.add_paragraph("")

        if not findings and not info_items:
            doc.add_paragraph("No security findings identified.")
            doc.save(output_path)
            return output_path

        for severity in ["High", "Medium", "Low"]:
            sev_findings = [f for f in findings if f.get("severity") == severity]
            if not sev_findings:
                continue

            doc.add_heading(f"{severity} ({len(sev_findings)})", 1)

            for f in sev_findings:
                doc.add_heading(f"{f['question_id']} — {f.get('finding_title', '')}", 2)

                p = doc.add_paragraph()
                p.add_run("Question: ").bold = True
                p.add_run(f.get("question", ""))

                p = doc.add_paragraph()
                run = p.add_run(f"Vendor answer: {f.get('vendor_answer', '')}")
                run.italic = True

                p = doc.add_paragraph()
                p.add_run("Finding: ").bold = True
                p.add_run(f.get("description", ""))

                p = doc.add_paragraph()
                p.add_run("Recommendation: ").bold = True
                p.add_run(f.get("recommendation", ""))

                doc.add_paragraph("")

        if info_items:
            doc.add_heading(f"Informational ({len(info_items)})", 1)
            for f in info_items:
                doc.add_heading(f"{f['question_id']} — {f.get('info_title', 'Document Referenced')}", 2)

                p = doc.add_paragraph()
                p.add_run("Question: ").bold = True
                p.add_run(f.get("question", ""))

                p = doc.add_paragraph()
                run = p.add_run(f"Vendor answer: {f.get('vendor_answer', '')}")
                run.italic = True

                p = doc.add_paragraph()
                p.add_run("Note: ").bold = True
                p.add_run("Vendor indicates a document has been provided. No finding raised.")

                p = doc.add_paragraph()
                p.add_run("Recommendation: ").bold = True
                p.add_run(f.get("recommendation", "Review the referenced document."))

                doc.add_paragraph("")

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # Output: markdown
    # ------------------------------------------------------------------

    def to_markdown(self) -> str:
        """Return findings as a markdown string grouped by severity."""
        findings = [f for f in self.findings if f.get("is_finding")]
        total = len(self.findings)
        counts = {s: sum(1 for f in findings if f.get("severity") == s) for s in ["High", "Medium", "Low"]}

        lines = [
            "# Completed Questionnaire — Security Findings\n",
            f"**Total questions reviewed:** {total} | "
            f"**Findings:** {len(findings)} | "
            f"Critical: {counts['Critical']} | High: {counts['High']} | "
            f"Medium: {counts['Medium']} | Low: {counts['Low']}\n",
        ]

        if not findings:
            lines.append("No security findings identified.")
            return "\n".join(lines)

        for severity in ["High", "Medium", "Low"]:
            sev_findings = [f for f in findings if f.get("severity") == severity]
            if not sev_findings:
                continue

            lines.append(f"## {severity} ({len(sev_findings)})\n")
            for f in sev_findings:
                lines.append(f"### {f['question_id']} — {f.get('finding_title', '')}\n")
                lines.append(f"**Question:** {f.get('question', '')}\n")
                lines.append(f"*Vendor answer: {f.get('vendor_answer', '')}*\n")
                lines.append(f"**Finding:** {f.get('description', '')}\n")
                lines.append(f"**Recommendation:** {f.get('recommendation', '')}\n")
                lines.append("---\n")

        return "\n".join(lines)
