"""
Security Document Reviewer - Extracts structured findings from SOC 2 reports
and penetration test reports.
"""
import os
import json
from typing import Dict, List, Any
from openai import OpenAI


class SecurityDocReviewer:

    COMPLETION_MODEL = "gpt-4o"

    def __init__(self):
        api_key = os.getenv("OPENAI_API_KEY")
        self.client = OpenAI(api_key=api_key) if api_key else None
        self.findings: List[Dict[str, Any]] = []

    def review(self, parsed_documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Identify and extract structured findings from each document."""
        self.findings = []

        if not self.client:
            return self.findings

        for doc in parsed_documents:
            filename = doc.get("filename", "unknown")
            full_text = self._get_text(doc)

            if not full_text.strip():
                continue

            doc_type = self._classify_document(filename, full_text)

            if doc_type == "soc":
                finding = self._extract_soc(filename, full_text)
            elif doc_type == "pentest":
                finding = self._extract_pentest(filename, full_text)
            else:
                finding = self._extract_other(filename, full_text)

            if finding:
                self.findings.append(finding)

        return self.findings

    # ------------------------------------------------------------------
    # Document classification
    # ------------------------------------------------------------------

    def _classify_document(self, filename: str, text: str) -> str:
        """Ask the LLM to classify the document as soc, pentest, or other."""
        sample = text[:3000]
        prompt = (
            "Classify the following document excerpt. Reply with exactly one word:\n"
            "- 'soc' if this is a SOC 2 or SOC 1 audit report\n"
            "- 'pentest' if this is a penetration test or security assessment report\n"
            "- 'other' for anything else\n\n"
            f"Filename: {filename}\n\nExcerpt:\n{sample}"
        )
        try:
            response = self.client.chat.completions.create(
                model=self.COMPLETION_MODEL,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=5,
                temperature=0
            )
            return response.choices[0].message.content.strip().lower()
        except Exception:
            return "other"

    # ------------------------------------------------------------------
    # SOC report extraction
    # ------------------------------------------------------------------

    def _extract_soc(self, filename: str, text: str) -> Dict[str, Any]:
        """Extract structured fields from a SOC report."""
        prompt = (
            "You are reviewing a SOC audit report. Extract the following fields from the document.\n"
            "If a field is not clearly stated, write 'Not specified'.\n\n"
            "Fields to extract:\n"
            "1. company_name: The name of the company being audited\n"
            "2. audit_period: The period covered by the audit (e.g. 'January 1, 2024 – December 31, 2024')\n"
            "3. auditor_company: The name of the auditing firm\n"
            "4. report_type: SOC type and report type (e.g. 'SOC 2 Type II')\n"
            "5. scope: Systems, services, and trust service criteria in scope\n"
            "6. auditors_opinion: The auditor's overall opinion (e.g. unqualified, qualified, adverse)\n"
            "7. exceptions: Any exceptions, deviations, or findings noted by the auditor. "
            "If none, write 'No exceptions noted'.\n\n"
            f"Document (filename: {filename}):\n{text[:15000]}\n\n"
            "Respond with a JSON object using exactly these keys: "
            "company_name, audit_period, auditor_company, report_type, scope, auditors_opinion, exceptions"
        )

        try:
            response = self.client.chat.completions.create(
                model=self.COMPLETION_MODEL,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
                temperature=0,
                response_format={"type": "json_object"}
            )
            data = json.loads(response.choices[0].message.content)
            data["document_type"] = "SOC Report"
            data["filename"] = filename
            return data
        except Exception as e:
            print(f"SOC extraction failed for {filename}: {e}")
            return {}

    # ------------------------------------------------------------------
    # Penetration test report extraction
    # ------------------------------------------------------------------

    def _extract_pentest(self, filename: str, text: str) -> Dict[str, Any]:
        """Extract structured fields from a penetration test report."""
        prompt = (
            "You are reviewing a penetration test or security assessment report. "
            "Extract the following fields from the document.\n"
            "If a field is not clearly stated, write 'Not specified'.\n\n"
            "Fields to extract:\n"
            "1. company_name: The name of the company that was tested\n"
            "2. testing_period: The dates or period during which testing was conducted\n"
            "3. pentester_company: The name of the firm that conducted the test\n"
            "4. scope: Systems, applications, or networks that were in scope\n"
            "5. findings: A JSON array of finding objects. Each object must have exactly two keys: "
            "\"severity\" (one of: Critical, High, Medium, Low, Informational) and "
            "\"description\" (a concise one-sentence summary of the finding). "
            "Example: [{\"severity\": \"High\", \"description\": \"SQL injection vulnerability in login endpoint\"}]. "
            "If no findings are listed, return an empty array [].\n"
            "6. remediation_status: A single string summarising the overall remediation status noted in the report "
            "(e.g. 'All findings remediated', 'In progress', 'Not provided'). "
            "If not noted, write 'Remediation status not provided'.\n\n"
            f"Document (filename: {filename}):\n{text[:15000]}\n\n"
            "Respond with a JSON object using exactly these keys: "
            "company_name, testing_period, pentester_company, scope, findings, remediation_status"
        )

        try:
            response = self.client.chat.completions.create(
                model=self.COMPLETION_MODEL,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1500,
                temperature=0,
                response_format={"type": "json_object"}
            )
            data = json.loads(response.choices[0].message.content)
            data["document_type"] = "Penetration Test Report"
            data["filename"] = filename
            return data
        except Exception as e:
            print(f"Pen test extraction failed for {filename}: {e}")
            return {}

    # ------------------------------------------------------------------
    # General document summary
    # ------------------------------------------------------------------

    def _extract_other(self, filename: str, text: str) -> Dict[str, Any]:
        """Extract a structured summary from any other document type (policies, guidelines, etc.)."""
        prompt = (
            "You are reviewing a vendor security or privacy document. "
            "Extract the following fields from the document.\n"
            "If a field is not clearly stated, write 'Not specified'.\n\n"
            "Fields to extract:\n"
            "1. document_title: The full title of the document\n"
            "2. document_category: The type of document "
            "(e.g. 'Information Security Policy', 'Privacy Policy', 'Data Retention Policy', "
            "'Acceptable Use Policy', 'Incident Response Policy', 'Business Continuity Plan', etc.)\n"
            "3. company_name: The name of the company that authored or owns the document\n"
            "4. effective_date: The date the document was approved, issued, or last reviewed "
            "(e.g. 'January 1, 2024')\n"
            "5. summary: A 2–4 sentence summary of what the document covers and its main purpose\n"
            "6. key_points: A JSON array of concise strings, each capturing one key point relevant "
            "to security or privacy (e.g. data handling practices, retention periods, access controls, "
            "encryption requirements, incident response steps). List up to 6 key points. "
            "If none are evident, return an empty array [].\n\n"
            f"Document (filename: {filename}):\n{text[:12000]}\n\n"
            "Respond with a JSON object using exactly these keys: "
            "document_title, document_category, company_name, effective_date, summary, key_points"
        )

        try:
            response = self.client.chat.completions.create(
                model=self.COMPLETION_MODEL,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=1000,
                temperature=0,
                response_format={"type": "json_object"}
            )
            data = json.loads(response.choices[0].message.content)
            data["document_type"] = "Policy / Other Document"
            data["filename"] = filename
            return data
        except Exception as e:
            print(f"Document summary extraction failed for {filename}: {e}")
            return {}

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _get_text(self, doc: Dict[str, Any]) -> str:
        """Reconstruct full text from a parsed document."""
        parts = []
        if doc["type"] == "pdf":
            for page in doc.get("pages", []):
                parts.append(page.get("text", ""))
        elif doc["type"] == "excel":
            for sheet in doc.get("sheets", []):
                for row in sheet.get("data", []):
                    parts.append(" ".join([str(c) for c in row if c]))
        return "\n".join(parts)

    def to_docx(self, output_path: str = "output/security_doc_review.docx") -> str:
        """Save structured findings to a Word document."""
        from docx import Document

        doc = Document()
        doc.add_heading("Security Document Review", 0)

        if not self.findings:
            doc.add_paragraph("No SOC or penetration test reports detected in the uploaded documents.")
        else:
            for f in self.findings:
                doc.add_heading(f"{f.get('document_type', 'Document')}: {f.get('filename', '')}", 1)

                if f.get("document_type") == "SOC Report":
                    for label, value in [
                        ("Company", f.get("company_name", "N/A")),
                        ("Report Type", f.get("report_type", "N/A")),
                        ("Audit Period", f.get("audit_period", "N/A")),
                        ("Auditing Firm", f.get("auditor_company", "N/A")),
                    ]:
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(value)
                    doc.add_heading("Scope", 2)
                    doc.add_paragraph(f.get("scope", "N/A"))
                    doc.add_heading("Auditor's Opinion", 2)
                    doc.add_paragraph(f.get("auditors_opinion", "N/A"))
                    doc.add_heading("Exceptions", 2)
                    doc.add_paragraph(f.get("exceptions", "N/A"))

                elif f.get("document_type") == "Penetration Test Report":
                    for label, value in [
                        ("Company", f.get("company_name", "N/A")),
                        ("Testing Period", f.get("testing_period", "N/A")),
                        ("Pen Tester", f.get("pentester_company", "N/A")),
                    ]:
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(value)
                    doc.add_heading("Scope", 2)
                    doc.add_paragraph(f.get("scope", "N/A"))
                    doc.add_heading("Findings", 2)
                    findings_list = f.get("findings", [])
                    if isinstance(findings_list, list) and findings_list and isinstance(findings_list[0], dict):
                        severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4}
                        findings_list = sorted(findings_list, key=lambda x: severity_order.get(x.get("severity", ""), 5))
                        tbl = doc.add_table(rows=1, cols=2)
                        tbl.style = "Table Grid"
                        hdr = tbl.rows[0].cells
                        hdr[0].paragraphs[0].add_run("Severity").bold = True
                        hdr[1].paragraphs[0].add_run("Description").bold = True
                        for item in findings_list:
                            row = tbl.add_row().cells
                            row[0].text = item.get("severity", "")
                            row[1].text = item.get("description", "")
                        doc.add_paragraph("")
                    elif isinstance(findings_list, list) and findings_list:
                        for item in findings_list:
                            doc.add_paragraph(str(item), style="List Bullet")
                    else:
                        doc.add_paragraph(str(findings_list) if findings_list else "No findings listed.")
                    doc.add_heading("Remediation Status", 2)
                    doc.add_paragraph(f.get("remediation_status", "N/A"))

                elif f.get("document_type") == "Policy / Other Document":
                    for label, value in [
                        ("Document Title", f.get("document_title", "N/A")),
                        ("Category", f.get("document_category", "N/A")),
                        ("Company", f.get("company_name", "N/A")),
                        ("Effective / Approval Date", f.get("effective_date", "N/A")),
                    ]:
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(value)
                    doc.add_heading("Summary", 2)
                    doc.add_paragraph(f.get("summary", "N/A"))
                    key_points = f.get("key_points", [])
                    if key_points:
                        doc.add_heading("Key Points", 2)
                        for point in key_points:
                            doc.add_paragraph(str(point), style="List Bullet")

                doc.add_paragraph("")

        doc.save(output_path)
        return output_path

    def to_markdown(self, output_path: str = "output/security_doc_review.md") -> str:
        """Save structured findings to a markdown file."""
        lines = ["# Security Document Review\n"]

        if not self.findings:
            lines.append("No SOC or penetration test reports detected in the uploaded documents.\n")
        else:
            for f in self.findings:
                lines.append(f"## {f.get('document_type', 'Document')}: {f.get('filename', '')}\n")

                if f.get("document_type") == "SOC Report":
                    lines.append(f"**Company:** {f.get('company_name', 'N/A')}")
                    lines.append(f"**Report Type:** {f.get('report_type', 'N/A')}")
                    lines.append(f"**Audit Period:** {f.get('audit_period', 'N/A')}")
                    lines.append(f"**Auditing Firm:** {f.get('auditor_company', 'N/A')}")
                    lines.append(f"\n**Scope:**\n{f.get('scope', 'N/A')}")
                    lines.append(f"\n**Auditor's Opinion:**\n{f.get('auditors_opinion', 'N/A')}")
                    lines.append(f"\n**Exceptions:**\n{f.get('exceptions', 'N/A')}")

                elif f.get("document_type") == "Penetration Test Report":
                    lines.append(f"**Company:** {f.get('company_name', 'N/A')}")
                    lines.append(f"**Testing Period:** {f.get('testing_period', 'N/A')}")
                    lines.append(f"**Pen Tester:** {f.get('pentester_company', 'N/A')}")
                    lines.append(f"\n**Scope:**\n{f.get('scope', 'N/A')}")
                    lines.append(f"\n**Findings:**\n{f.get('findings', 'N/A')}")
                    lines.append(f"\n**Remediation Status:**\n{f.get('remediation_status', 'N/A')}")

                elif f.get("document_type") == "Policy / Other Document":
                    lines.append(f"**Document Title:** {f.get('document_title', 'N/A')}")
                    lines.append(f"**Category:** {f.get('document_category', 'N/A')}")
                    lines.append(f"**Company:** {f.get('company_name', 'N/A')}")
                    lines.append(f"**Effective / Approval Date:** {f.get('effective_date', 'N/A')}")
                    lines.append(f"\n**Summary:**\n{f.get('summary', 'N/A')}")
                    key_points = f.get("key_points", [])
                    if key_points:
                        lines.append("\n**Key Points:**")
                        for point in key_points:
                            lines.append(f"- {point}")

                lines.append("\n---\n")

        content = "\n".join(lines)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        return output_path
